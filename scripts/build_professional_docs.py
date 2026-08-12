from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx"
OUT.mkdir(parents=True, exist_ok=True)

BLUE = "3157D5"
PURPLE = "7651C9"
INK = "172033"
MUTED = "667085"
LIGHT = "F4F6FC"
LAVENDER = "EEEAFB"
WHITE = "FFFFFF"
BORDER = "D7DCEC"
TODAY = "12 de agosto de 2026"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=10.5, color=INK, bold=False, italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def configure_document(title):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(9.8)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Heading 1", 16.5, BLUE, 12, 6),
        ("Heading 2", 13.2, PURPLE, 9, 5),
        ("Heading 3", 11.5, INK, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("RUAHTONE  /  DOCUMENTACIÓN OFICIAL")
    set_font(r, 8.5, MUTED, True)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{title}  •  {TODAY}")
    set_font(r, 8, MUTED)
    return doc


def add_cover(doc, kicker, title, subtitle, code, version="Versión 1.0"):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run(kicker.upper())
    set_font(r, 10, PURPLE, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_font(r, 30, BLUE, True, name="Aptos Display")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run(subtitle)
    set_font(r, 15, INK)
    add_callout(doc, "PROPÓSITO", "Fuente profesional para diseñar, construir y validar el MVP frontend local de RUAHTONE.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    for text, bold in ((f"{code}  •  {version}", True), (f"\n{TODAY}", False), ("\nEstado: Documento oficial del MVP", False)):
        r = p.add_run(text)
        set_font(r, 10, MUTED, bold)
    doc.add_page_break()


def add_intro(doc, source_note=True):
    if source_note:
        add_callout(doc, "CONTROL DE FUENTE", "Este documento reorganiza exclusivamente la Instrucción Principal RUAHTONE — Master Product & Development Prompt. No agrega funcionalidades al alcance oficial.")


def add_callout(doc, label, text, color=BLUE, trailing_space=True):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_font(r, 8.5, color, True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, 10.5, INK)
    if trailing_space:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.34)
        p.paragraph_format.first_line_indent = Inches(-0.17)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_font(r)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    if widths is None:
        base = 9360 // len(headers)
        widths = [base] * len(headers)
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, 9, WHITE, True)
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if row_idx % 2:
                set_cell_shading(cells[i], "F8F9FC")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_font(r, 9.2)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_section(doc, heading, paragraphs=None, bullets=None):
    doc.add_heading(heading, level=1)
    for text in paragraphs or []:
        p = doc.add_paragraph(text)
        p.paragraph_format.widow_control = True
    if bullets:
        add_bullets(doc, bullets)


def save(doc, filename, title, subject):
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "RUAHTONE"
    props.keywords = "RUAHTONE, iglesia, servicios, MVP, producto"
    props.comments = "Documento generado desde la especificación oficial de RUAHTONE."
    doc.save(OUT / filename)


def build_master():
    title = "Documento Maestro del Producto"
    doc = configure_document(title)
    add_cover(doc, "RUAHTONE", title, "Visión, filosofía, alcance y experiencia del MVP", "RHT-PROD-001")
    add_intro(doc)
    add_section(doc, "1. Resumen ejecutivo", [
        "RUAHTONE es un sistema operativo para la organización de una iglesia y sus ministerios. Centraliza la preparación, coordinación y ejecución de servicios para evitar que WhatsApp funcione como agenda, repositorio, lista de asistencia e historial improvisado.",
        "El MVP será un frontend local, realista y funcional, pensado para probar personalmente, validar con miembros reales de una iglesia, descubrir errores de experiencia y demostrar el producto antes de convertirlo en un SaaS.",
    ])
    add_callout(doc, "PROPUESTA DE VALOR", "Organizar un servicio de iglesia sin depender de WhatsApp.", PURPLE)
    add_section(doc, "2. Problema observado", ["La idea nace de la experiencia directa del fundador en adoración, guitarra, sonido y coordinación de servicios en una iglesia de Chile."], [
        "Horarios, confirmaciones, tonalidades, PDFs, canciones, ensayos y cambios se dispersan entre mensajes.",
        "La información se pierde, se repite y queda enterrada.",
        "La preparación depende demasiado de una persona que organiza.",
        "Existe poca visibilidad del estado actual y dificultad para saber qué cambió.",
        "Los servicios especiales aumentan la complejidad técnica y operacional.",
    ])
    add_section(doc, "3. Principio central", ["El servicio es la unidad operacional principal. Las pantallas y entidades existen en función de la preparación real del servicio, no como módulos administrativos desconectados."])
    add_table(doc, ["Nivel", "Relación operacional"], [
        ("Iglesia", "Organiza ministerios y personas"),
        ("Ministerios", "Participan en servicios"),
        ("Servicios", "Concentran la preparación"),
        ("Asignaciones", "Relacionan personas, ministerios y funciones"),
        ("Canciones y recursos", "Definen qué debe prepararse"),
        ("Checklists y actividad", "Muestran ejecución, cambios e historial"),
    ], [1900, 7460])
    add_section(doc, "4. Usuarios y experiencia emocional")
    add_table(doc, ["Usuario", "Debe responder rápidamente", "Experiencia buscada"], [
        ("Pastor", "¿Está preparado? ¿Hay problemas?", "Tranquilidad"),
        ("Líder", "¿Quién confirmó? ¿Qué falta?", "Control"),
        ("Músico / servidor", "¿Cuándo, dónde y qué preparo?", "Claridad"),
        ("Sonido", "¿Qué configuración y tareas necesito?", "Preparación"),
        ("Multimedia", "¿Qué presentación y recursos están listos?", "Orden"),
        ("Administrador", "¿Cómo está la operación general?", "Orden"),
    ], [1600, 5000, 2760])
    add_section(doc, "5. Objetivo del MVP", bullets=[
        "Crear y configurar un servicio.", "Asignar ministerios, personas y funciones específicas.",
        "Permitir que cada servidor confirme su propia asistencia.", "Agregar canciones y definir tonalidad por servicio.",
        "Asociar recursos a entidades del dominio.", "Crear checklists contextuales.",
        "Visualizar cambios e información del servicio en un solo lugar.",
    ])
    add_section(doc, "6. Alcance oficial", ["El MVP funciona completamente en local con datos demo realistas de Iglesia Manantiales, Santiago de Chile."], [
        "Inicio contextual, servicios, Centro del Servicio, calendario y búsqueda global.",
        "Personas, roles, ministerios, asignaciones y confirmaciones.",
        "Canciones, instancias por servicio, recursos e historial de uso.",
        "Checklists de servicio, ministerio y persona; actividad estructurada y notificaciones relevantes.",
        "Contextos específicos para servidor, líder, administrador, pastor, sonido y multimedia.",
        "Sonido con monitoreo, sala, P16, patch, escenas y prueba de sonido como información operacional.",
        "Temas claro y oscuro, sidebar adaptable, responsive y restablecimiento de datos demo.",
    ])
    add_section(doc, "7. Fuera de alcance", bullets=[
        "Backend, base de datos remota y autenticación real.", "Pagos, suscripciones y facturación.",
        "Aplicaciones iOS, Android, React Native o Capacitor.", "Chat, API de WhatsApp y funciones sociales.",
        "IA, reconocimiento de audio e integraciones complejas.", "Contabilidad, diezmos y ofrendas.",
        "Streaming, sincronización cloud y control físico de consolas.",
    ])
    add_section(doc, "8. Filosofía de diseño", ["La interfaz debe sentirse premium porque es simple: organizada, clara y tranquila; cristiana sin clichés visuales, profesional sin ser corporativa y moderna sin ser fría."], [
        "Referencias de sensación: Apple, Linear, Notion, Stripe y Superhuman, sin copiar interfaces.",
        "Base visual minimalista con azul y morado como acentos controlados.",
        "Light y Dark Mode con jerarquía, contraste y legibilidad propios.",
        "Tipografía moderna, espacio generoso, bordes suaves y sombras discretas.",
        "Animaciones sutiles solo donde transmiten estado o continuidad.",
        "No usar gradientes, glassmorphism, color o movimiento de forma excesiva.",
    ])
    add_section(doc, "9. Principios de experiencia", bullets=[
        "El usuario debe entender lo esencial en menos de diez segundos.",
        "Contexto, estado y acción tienen prioridad sobre información secundaria.",
        "Cada pantalla tiene una acción principal clara.",
        "No existen pantallas vacías sin explicación ni botones principales sin comportamiento.",
        "La navegación y los permisos se adaptan al contexto real del usuario.",
        "El calendario permite encontrar servicios; no es el centro del producto.",
        "La actividad es historial estructurado, no mensajería.",
    ])
    add_section(doc, "10. Criterios absolutos", bullets=[
        "Ante conflicto entre estética y funcionamiento, gana el funcionamiento.",
        "Ante conflicto entre cantidad y simplicidad, gana la simplicidad.",
        "Ante conflicto entre un modelo genérico y la operación real de una iglesia, gana la operación real.",
        "El MVP demuestra una cosa excepcionalmente bien: preparar y coordinar un servicio sin el caos de WhatsApp.",
    ])
    add_callout(doc, "CRITERIO DE ÉXITO", "Un miembro entiende qué debe hacer; un líder deja de reenviar todo por WhatsApp; y la iglesia percibe que la preparación está ordenada.", trailing_space=False)
    save(doc, "01_RUAHTONE_Documento_Maestro_Producto.docx", title, "Visión y alcance oficial del MVP")


def build_functional():
    title = "Especificación Funcional del MVP"
    doc = configure_document(title)
    add_cover(doc, "RUAHTONE", title, "Requisitos, reglas de negocio, estados y flujos obligatorios", "RHT-FUNC-001")
    add_intro(doc)
    add_section(doc, "1. Regla de oro", ["Antes de implementar una función debe responderse qué problema operativo real de una iglesia resuelve. Si la respuesta no es clara, la función no debe inventarse."])
    add_section(doc, "2. Servicios", bullets=[
        "Crear, editar, duplicar, eliminar, publicar y archivar.",
        "Registrar nombre, tipo, fecha, hora de inicio, hora de término, lugar y descripción.",
        "Concentrar ministerios, personas, canciones, recursos, checklist y actividad.",
        "Estados: Planificación, Organizando, Confirmando, Listo, Completado y Archivado.",
        "Tipos: Culto, Ensayo, Conferencia, Jóvenes, Reunión y Evento especial.",
        "Un ensayo puede relacionarse con el servicio o ministerio que prepara.",
        "Duplicar conserva estructura seleccionada y reinicia confirmaciones.",
    ])
    add_section(doc, "3. Personas, roles y ministerios", ["Una persona posee una identidad única y puede servir en varios ministerios. Rol y ministerio nunca son equivalentes."])
    add_table(doc, ["Concepto", "Definición oficial", "Ejemplo"], [
        ("Persona", "Identidad única dentro de la iglesia", "Benjamín"),
        ("Rol", "Responsabilidad general", "Servidor / Líder"),
        ("Ministerio", "Área donde la persona sirve", "Adoración, Sonido"),
        ("Asignación", "Participación contextual en un servicio", "Primera guitarra"),
    ], [1600, 4560, 3200])
    add_section(doc, "4. Asignaciones y asistencia", bullets=[
        "La asignación relaciona servicio, persona, ministerio y función específica.",
        "Funciones posibles incluyen instrumentos, voces, sala, monitores, multimedia y ujieres.",
        "El servidor confirma su propia asistencia; líder y administrador observan el estado.",
        "Estados: Pendiente, Confirmado, No asistiré y Llegaré tarde.",
        "Ausencia y atraso producen alertas visibles.",
        "Una doble función incompatible en el mismo horario produce advertencia, sin bloqueo obligatorio en el MVP.",
    ])
    add_section(doc, "5. Canciones y tonalidad contextual", ["La canción conserva su tonalidad original. Cada uso dentro de un servicio posee configuración propia; cambiarla no modifica la biblioteca."])
    add_table(doc, ["Entidad", "Datos relevantes"], [
        ("Song", "Nombre, autor, tonalidad original, BPM, duración, compás, etiquetas, recursos e historial de uso"),
        ("ServiceSong", "Servicio, canción, orden, tonalidad del servicio, BPM opcional, versión, notas y estado"),
    ], [2200, 7160])
    add_bullets(doc, [
        "El líder puede reordenar mediante drag and drop y botones arriba/abajo como alternativa.",
        "El detalle de canción muestra información general, recursos, historial y servicios donde fue utilizada.",
        "Los contadores y valores visibles deben derivarse de datos reales y consistentes.",
    ])
    add_section(doc, "6. Recursos", ["Los recursos no son una carpeta aislada: deben asociarse a una canción, servicio o ministerio."], [
        "Tipos: PDF, Partitura, Chord Chart, Secuencia, Multitrack, Audio, Video, Imagen, Documento y Otro.",
        "Cada recurso debe ofrecer una acción visual coherente.",
        "Los archivos simulados deben presentarse explícitamente como parte del MVP local.",
    ])
    add_section(doc, "7. Checklists, actividad y notificaciones", bullets=[
        "Los checklists son contextuales; un servicio normal y uno especial pueden requerir estructuras distintas.",
        "Cada ministerio puede tener su checklist y cada servidor tareas personales.",
        "Completar tareas actualiza el estado visible.",
        "Cambios relevantes generan actividad estructurada con actor y momento.",
        "La actividad no incluye likes, reacciones, grupos ni conversaciones.",
        "Las notificaciones deben ser relevantes y no generar spam.",
    ])
    add_section(doc, "8. Centro del Servicio", ["Es la pantalla más importante del MVP y el destino común desde Inicio, Calendario, Servicios, notificaciones y búsqueda."], [
        "Fecha, horario, lugar y estado.", "Ministerios y contadores de confirmación.",
        "Personas, funciones y asistencia.", "Canciones y tonalidades del servicio.",
        "Recursos, checklist, actividad y observaciones.", "Acción principal: Gestionar servicio.",
    ])
    add_section(doc, "9. Experiencias por contexto")
    add_table(doc, ["Contexto", "Prioridad de inicio", "Capacidad principal"], [
        ("Músico", "Próximo servicio, confirmación, canciones, cambios, ensayo", "Responder asistencia y preparar repertorio"),
        ("Líder", "Servicio, confirmaciones, pendientes, canciones, tareas", "Gestionar su ministerio en el servicio"),
        ("Administrador", "Servicios, ministerios, personas, estado general", "Administrar estructura y operación"),
        ("Pastor", "Servicios, preparación, ministerios, alertas", "Ver estado sin configuración técnica"),
        ("Sonido", "Servicio, checklist, configuración, P16 y patch", "Preparar operación técnica"),
    ], [1400, 4400, 3560])
    add_section(doc, "10. Sonido", bullets=[
        "Separar sala, monitoreo y extensión conceptual para streaming.",
        "Representar integrantes, instrumentos, canales, sistemas de monitoreo y observaciones.",
        "Permitir configuraciones P16 diferentes por servicio.",
        "Registrar patch con canal, fuente, micrófono, instrumento, entrada, salida y observación.",
        "Asociar una escena a un servicio.",
        "Registrar prueba con o sin secuencia y estado Pendiente, En progreso o Completada.",
        "Mostrar cambios respecto del servicio anterior.",
    ])
    add_section(doc, "11. Navegación y configuración", bullets=[
        "Rutas: Inicio, Servicios, Centro del Servicio, Calendario, Canciones, Personas, Recursos, Ministerios y Configuración.",
        "Búsqueda global mediante Cmd + K o Ctrl + K para servicios, canciones, personas, recursos y ministerios.",
        "Sidebar expandible y colapsable; navegación compacta en móvil.",
        "Configuración: tema, información de iglesia, demo, preferencias y restablecimiento confirmado.",
        "El cambio de contexto modifica navegación, dashboard, información y acciones.",
    ])
    add_section(doc, "12. Flujos obligatorios de demostración")
    add_numbered(doc, [
        "Flujo principal: próximo servicio, Centro, asignaciones, confirmación, contexto líder, cambio de tono contextual, recurso, actividad, checklist, calendario y duplicación con confirmaciones pendientes.",
        "Sonido: servicio, checklist, músicos, P16, cambios, patch y preparación completada.",
        "Líder de adoración: confirmaciones, agregar y ordenar canción, cambiar tono, nota, PDF y actividad.",
        "Pastor: estado general, ministerios, alertas y Centro del Servicio sin configuración técnica.",
    ])
    add_section(doc, "13. Criterios de aceptación global", bullets=[
        "Todas las rutas funcionan y la aplicación compila sin errores TypeScript.",
        "LocalStorage persiste creación, edición, eliminación, duplicación y cambios.",
        "Confirmaciones, tonalidades, recursos, checklist y actividad actualizan solo las entidades correspondientes.",
        "Búsqueda, tema, responsive y estados vacíos funcionan.",
        "Los datos demo son realistas y sus estadísticas coinciden.",
        "Las acciones principales son reales o se simulan de forma coherente y explícita.",
    ])
    save(doc, "02_RUAHTONE_Especificacion_Funcional_MVP.docx", title, "Requisitos funcionales oficiales")


def build_domain():
    title = "Modelo de Dominio y Arquitectura"
    doc = configure_document(title)
    add_cover(doc, "RUAHTONE", title, "Entidades, relaciones, fronteras de datos y tecnología oficial", "RHT-ARCH-001")
    add_intro(doc)
    add_section(doc, "1. Arquitectura conceptual", ["Los datos se modelan relacionalmente aunque el MVP no utilice backend. Las relaciones guardan identificadores y evitan duplicar información innecesaria."])
    add_table(doc, ["Entidad", "Responsabilidad oficial"], [
        ("Church", "Organización principal; futura relación con plan y usuarios"),
        ("User", "Identidad de acceso conceptual, separada de permisos y ministerios"),
        ("Person", "Miembro o participante con rol, ministerios, disponibilidad y estado"),
        ("Ministry", "Área de servicio con líder, miembros, checklist y configuraciones"),
        ("Service", "Unidad operacional central"),
        ("ServiceAssignment", "Persona asignada a función y ministerio en un servicio"),
        ("Song", "Información global de biblioteca y tonalidad original"),
        ("ServiceSong", "Configuración contextual de una canción en un servicio"),
        ("Resource", "Recurso relacionado con canción, servicio o ministerio"),
        ("Checklist / Item", "Preparación contextual y estado de tareas"),
        ("Activity", "Historial estructurado de cambios"),
        ("Notification", "Aviso relevante para el usuario"),
        ("SoundConfiguration", "Preparación técnica de sala, monitoreo, P16, patch y escena"),
    ], [2500, 6860])
    add_section(doc, "2. Relaciones principales")
    add_numbered(doc, [
        "Una Iglesia organiza Ministerios, Personas y Servicios.",
        "Una Persona puede pertenecer a varios Ministerios.",
        "Un Servicio contiene Asignaciones que referencian Persona y Ministerio.",
        "Un Servicio contiene ServiceSongs que referencian una Song global.",
        "Recursos se relacionan con Canciones, Servicios o Ministerios.",
        "Checklists se contextualizan por Servicio, Ministerio o Persona.",
        "La Actividad registra modificaciones relevantes del Servicio.",
    ])
    add_section(doc, "3. Diferencia entre datos globales y contextuales")
    add_table(doc, ["Dato global", "Dato contextual"], [
        ("Song.originalKey = D", "ServiceSong.key = F para Culto Domingo"),
        ("Person y sus ministerios", "ServiceAssignment y función en un servicio"),
        ("Recurso de una canción", "Recurso asociado a un servicio específico"),
        ("Plantilla de checklist", "Checklist ejecutado para un servicio"),
    ], [4680, 4680])
    add_callout(doc, "INVARIANTE CRÍTICA", "Cambiar la tonalidad de una canción dentro de un servicio modifica únicamente ServiceSong. Song.originalKey permanece intacta.", PURPLE)
    add_section(doc, "4. Contratos mínimos oficiales")
    add_table(doc, ["Entidad", "Campos mínimos definidos"], [
        ("Person", "id, nombre, apellido, avatar, teléfono opcional, ministerios, rol, disponibilidad, estado"),
        ("ServiceAssignment", "id, serviceId, personId, ministryId, función, attendanceStatus, notas"),
        ("Song", "id, nombre, autor, tonalidad original, BPM, duración, compás, etiquetas, recursos, historial"),
        ("ServiceSong", "id, serviceId, songId, orden, key, BPM opcional, versión, notas, enabled"),
        ("Service", "nombre, fecha, horas, lugar, descripción, ministerios, personas, canciones, recursos, checklist, actividad"),
        ("Ministry", "id, nombre, descripción, icono, color, líder, miembros, checklist, configuraciones"),
    ], [2300, 7060])
    add_section(doc, "5. Reglas de integridad", bullets=[
        "Las relaciones almacenan IDs, no nombres duplicados.",
        "Las estadísticas se calculan desde asignaciones y estados reales.",
        "Una confirmación cambia únicamente la asignación de la persona que responde.",
        "Una persona puede tener varios ministerios, pero una sola identidad.",
        "Los cambios importantes generan actividad.",
        "Los servicios históricos conservan canciones, tonalidades y asignaciones utilizadas.",
        "Un conflicto de horario se informa claramente.",
        "Los permisos se simulan aunque no exista autenticación real.",
    ])
    add_section(doc, "6. Estados oficiales")
    add_table(doc, ["Entidad", "Estados"], [
        ("Servicio", "Planificación, Organizando, Confirmando, Listo, Completado, Archivado"),
        ("Asistencia", "Pendiente, Confirmado, No asistiré, Llegaré tarde"),
        ("Prueba de sonido", "Pendiente, En progreso, Completada"),
        ("Checklist", "Ítems completos o pendientes"),
    ], [2600, 6760])
    add_section(doc, "7. Tecnología oficial del MVP", bullets=[
        "React, TypeScript y Vite.", "Tailwind CSS y shadcn/ui.", "Lucide Icons y Framer Motion.",
        "LocalStorage como persistencia local.", "React Context o Zustand solamente si simplifica el estado.",
        "Tipos e interfaces TypeScript para el dominio.",
    ])
    add_section(doc, "8. Restricciones técnicas", bullets=[
        "No utilizar Firebase, Supabase, PostgreSQL, MongoDB ni APIs externas.",
        "No implementar backend, autenticación real ni servicios cloud.",
        "Crear un store central para personas, ministerios, servicios, asignaciones, canciones, recursos, checklists, actividad y configuración.",
        "Persistir cambios en LocalStorage y permitir restablecer el seed demo.",
        "Separar conceptualmente usuario, permisos, rol, ministerios y datos.",
    ])
    add_section(doc, "9. Preparación para evolución futura", ["Sin implementar estas capacidades, el diseño no debe bloquear aplicación móvil, autenticación, múltiples iglesias, multi-sede, pagos, notificaciones, almacenamiento cloud, roles personalizados, estadísticas, IA e integraciones."], [
        "La futura estructura SaaS contempla Iglesia → Plan → Usuarios → Ministerios → Servicios.",
        "La lógica conceptual debe poder reutilizarse en una experiencia móvil.",
        "Sonido contempla streaming como área futura, sin activarlo en el MVP.",
    ])
    add_section(doc, "10. Estructura lógica sugerida por la especificación", ["Esta vista organiza las prioridades oficiales; no introduce una dependencia adicional."], [
        "Capa de dominio: entidades, relaciones, estados y reglas.",
        "Store central: acciones y persistencia LocalStorage.",
        "Experiencias por contexto: servidor, líder, administrador, pastor, sonido y multimedia.",
        "Features operacionales: servicios, asignaciones, canciones, recursos, checklists, actividad, calendario y sonido.",
        "Datos demo: Iglesia Manantiales y escenarios consistentes.",
    ])
    save(doc, "03_RUAHTONE_Modelo_Dominio_Arquitectura.docx", title, "Modelo y arquitectura oficial del MVP")


def build_roadmap():
    title = "Roadmap y Plan de Validación"
    doc = configure_document(title)
    add_cover(doc, "RUAHTONE", title, "Secuencia oficial de implementación, demostración y calidad", "RHT-PLAN-001")
    add_intro(doc)
    add_section(doc, "1. Estrategia de entrega", ["La prioridad oficial es calidad sobre cantidad. Es preferible entregar seis pantallas excelentes que veinte mediocres. Cada incremento debe sostener una experiencia de producto, no un CRUD aislado."])
    add_section(doc, "2. Orden oficial de implementación")
    priorities = [
        "Arquitectura de datos", "Layout global", "Inicio contextual", "Servicios", "Centro del Servicio",
        "Asignaciones", "Confirmaciones", "Canciones", "Tonalidades por servicio", "Recursos",
        "Checklists", "Actividad", "Calendario", "Personas", "Ministerios", "Sonido", "Configuración",
    ]
    add_table(doc, ["Prioridad", "Área"], [(i + 1, value) for i, value in enumerate(priorities)], [1400, 7960])
    add_section(doc, "3. Hitos de trabajo", ["Los hitos siguientes agrupan el orden oficial sin cambiarlo."])
    add_table(doc, ["Hito", "Contenido", "Resultado verificable"], [
        ("Fundación", "Datos, layout e Inicio", "Base local consistente y experiencia contextual"),
        ("Núcleo", "Servicios, Centro, asignaciones y confirmaciones", "Coordinación básica del servicio"),
        ("Preparación", "Canciones, tonalidades, recursos, checklist y actividad", "Información musical y operacional centralizada"),
        ("Descubrimiento", "Calendario, personas y ministerios", "Acceso y administración de entidades relacionadas"),
        ("Especialización", "Sonido y configuración", "Flujo técnico y demo reiniciable"),
        ("Validación", "Testing integral y pruebas reales", "MVP listo para evaluación"),
    ], [1450, 4200, 3710])
    add_section(doc, "4. Demo flow principal")
    add_numbered(doc, [
        "Abrir RUAHTONE y ver el próximo servicio.", "Abrir el Centro del Servicio y ver personas asignadas.",
        "Confirmar asistencia como servidor.", "Volver al contexto de líder y ver la actualización.",
        "Abrir una canción y ver su tonalidad original.", "Cambiar la tonalidad solo para el servicio y verificar que el original no cambió.",
        "Agregar un recurso y revisar la actividad.", "Completar una tarea y ver el cambio en dashboard.",
        "Abrir el calendario y regresar al servicio.", "Duplicar el servicio, cambiar fecha y verificar confirmaciones pendientes.",
    ])
    add_section(doc, "5. Demos por rol")
    add_table(doc, ["Rol", "Recorrido obligatorio"], [
        ("Sonido", "Servicio → checklist → músicos → P16 → cambios → patch → preparación completada"),
        ("Líder de adoración", "Confirmaciones → agregar canción → ordenar → tono → nota → PDF → actividad"),
        ("Pastor", "Inicio → estado general → ministerios → alertas → Centro del Servicio"),
    ], [2200, 7160])
    add_section(doc, "6. Plan de pruebas oficial", bullets=[
        "Navegación y rutas.", "LocalStorage.", "Creación, edición, eliminación y duplicación.",
        "Confirmaciones y cambio de tonalidad.", "Recursos, checklist y actividad.",
        "Búsqueda global y tema.", "Responsive y estados vacíos.",
        "Consistencia de contadores y datos demo.", "Permisos y cambio real de experiencia por contexto.",
    ])
    add_section(doc, "7. Checklist de entrega")
    add_table(doc, ["Control", "Condición de salida"], [
        ("Compilación", "El proyecto compila"), ("TypeScript", "Sin errores"),
        ("Consola", "Sin errores importantes"), ("Rutas", "Todas funcionan"),
        ("Persistencia", "LocalStorage conserva cambios correctamente"), ("Demo", "Datos completos y flujos funcionales"),
        ("Interfaz", "Responsive con Light y Dark Mode"), ("Interacciones", "Acciones principales reales"),
        ("Entrega", "Arquitectura, estructura, funciones, simulaciones, decisiones e instrucciones documentadas"),
    ], [2500, 6860])
    add_section(doc, "8. Validación con usuarios", ["El MVP se utiliza para pruebas personales y con miembros reales de una iglesia, descubrir errores de UX, validar si resuelve los problemas y demostrar el producto."], [
        "Servidor: identificar servicio, horario, lugar, confirmación, canciones y cambios en menos de diez segundos.",
        "Líder: identificar confirmados, pendientes, canciones, recursos faltantes y tareas.",
        "Sonido: identificar configuración, cambios, músicos, P16 y checklist.",
        "Pastor: identificar preparación, problemas, personas faltantes y cambios importantes.",
    ])
    add_section(doc, "9. Mejoras futuras excluidas del MVP", bullets=[
        "Aplicaciones móviles y notificaciones push.", "Autenticación, múltiples iglesias y multi-sede.",
        "Suscripciones, pagos y almacenamiento cloud.", "Roles personalizados y estadísticas.",
        "IA, reemplazos inteligentes e integraciones musicales o de sonido.", "Streaming e integraciones adicionales.",
    ])
    add_callout(doc, "REGLA DE CIERRE", "El MVP no debe intentar demostrar que RUAHTONE puede hacerlo todo. Debe demostrar que prepara y coordina un servicio excepcionalmente bien.", trailing_space=False)
    save(doc, "04_RUAHTONE_Roadmap_Plan_Validacion.docx", title, "Roadmap oficial y plan de validación")


def build_index():
    title = "Índice Ejecutivo de Documentación"
    doc = configure_document(title)
    add_cover(doc, "RUAHTONE", title, "Mapa del paquete documental oficial del MVP", "RHT-DOC-000")
    add_intro(doc)
    add_section(doc, "1. Propósito del paquete", ["Transformar la instrucción principal en documentos profesionales, utilizables y coherentes sin alterar la idea oficial ni agregar alcance."])
    add_section(doc, "2. Documentos incluidos")
    add_table(doc, ["Código", "Documento", "Uso principal"], [
        ("RHT-PROD-001", "Documento Maestro del Producto", "Alinear visión, problema, alcance y experiencia"),
        ("RHT-FUNC-001", "Especificación Funcional del MVP", "Diseñar y aceptar comportamiento"),
        ("RHT-ARCH-001", "Modelo de Dominio y Arquitectura", "Implementar datos y fronteras correctamente"),
        ("RHT-PLAN-001", "Roadmap y Plan de Validación", "Secuenciar, probar y entregar el MVP"),
    ], [1900, 3600, 3860])
    add_section(doc, "3. Orden de lectura", bullets=[
        "Fundadores y stakeholders: Documento Maestro → Roadmap.",
        "Diseño de producto: Documento Maestro → Especificación Funcional → Roadmap.",
        "Desarrollo: Especificación Funcional → Modelo de Dominio → Roadmap.",
        "Pruebas y validación: Especificación Funcional → Roadmap.",
    ])
    add_section(doc, "4. Control documental")
    add_table(doc, ["Campo", "Valor"], [
        ("Fuente", "RUAHTONE — Master Product & Development Prompt"),
        ("Versión del paquete", "1.0"), ("Fecha", TODAY),
        ("Estado", "Oficial para definición del MVP"),
        ("Principio de cambio", "Toda modificación de alcance debe volver a contrastarse con la idea oficial"),
    ], [2200, 7160])
    add_callout(doc, "IDENTIDAD", "RUAHTONE — Todo lo que necesitas para preparar el servicio. En un solo lugar.", PURPLE, trailing_space=False)
    save(doc, "00_RUAHTONE_Indice_Ejecutivo.docx", title, "Índice del paquete documental")


if __name__ == "__main__":
    build_index()
    build_master()
    build_functional()
    build_domain()
    build_roadmap()
    print(f"Generated 5 DOCX files in {OUT}")
