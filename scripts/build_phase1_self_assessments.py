from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PHASE = ROOT / "APT122_005D" / "Equipo_Por_Definir" / "Fase_1"
OUTPUT_ROOT = PHASE / "Evidencias_Individuales"

BLUE = "3157D5"
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "667085"
LIGHT_BLUE = "EEF3FF"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "D7DCEC"
GREEN = "247A52"
GOLD = "8A6500"

COURSE = "Capstone"
COURSE_CODE = "PTY4614"
SECTION = "005D"
PROJECT = "RUAHTONE"
DATE_TEXT = "12 de agosto de 2026"


STUDENTS = [
    {
        "first": "Benjamín",
        "last": "Olmedo",
        "folder": "Olmedo_Benjamin",
        "filename": "Olmedo_Benjamin_1.3_APT122_AutoevaluacionFase1.docx",
        "focus_es": (
            "Mis intereses profesionales se concentran en el desarrollo de productos digitales útiles, la arquitectura de software y la mejora de procesos mediante tecnología. "
            "Además, mi experiencia sirviendo en sonido y adoración me permite reconocer el problema desde la operación real de una iglesia. RUAHTONE integra esos intereses porque exige convertir una necesidad observada en un producto organizado, usable, escalable y verificable."
        ),
        "role_es": (
            "Aporto principalmente la visión de sonido y adoración, el levantamiento de reglas operacionales, la coherencia del modelo de dominio y la coordinación de la documentación del producto."
        ),
        "conclusion_en": (
            "RUAHTONE is a relevant and feasible Capstone project because it addresses a recurring coordination problem with a clearly bounded local MVP. "
            "The current proposal connects user needs with project management, data modeling, software development, and validation practices. "
            "Its strongest point is the service-centered model, which keeps assignments, songs, resources, checklists, and activity in one coherent operational space. "
            "The next stage must transform the approved specifications into a tested vertical prototype and collect evidence from real church volunteers."
        ),
        "reflection_en": (
            "This self-assessment helped me separate the long-term business vision from the product that we can responsibly build during the course. "
            "My practical experience in sound and worship is valuable, but it must be translated into explicit requirements, acceptance criteria, and testable decisions. "
            "I also identified that our technical ambition can become a risk if we expand the scope too early. "
            "I will contribute by protecting the agreed scope, documenting domain rules, coordinating evidence, and validating the sound and worship workflows with real users."
        ),
    },
    {
        "first": "Daniel",
        "last": "Baeza",
        "folder": "Baeza_Daniel",
        "filename": "Baeza_Daniel_1.3_APT122_AutoevaluacionFase1.docx",
        "focus_es": (
            "Mis intereses profesionales se relacionan con el desarrollo de software, la experiencia de usuario y la creación de herramientas que ordenen procesos cotidianos. "
            "Mi participación en multimedia dentro de una iglesia me permite observar retrasos en letras, presentaciones, recursos y cambios que llegan por canales dispersos. "
            "RUAHTONE conecta esa experiencia con el diseño de una interfaz clara y con la implementación sistemática de funcionalidades que deben ser comprensibles para distintos roles."
        ),
        "role_es": (
            "Aporto la perspectiva operativa de multimedia, la revisión de los flujos de recursos y contenidos, y la evaluación de claridad, accesibilidad y respuesta de la interfaz."
        ),
        "conclusion_en": (
            "RUAHTONE provides a coherent answer to an observable communication and preparation problem in churches. "
            "The proposal is relevant to software development because it requires structured requirements, role-based experiences, persistent data, accessible interfaces, and systematic testing. "
            "The local MVP is achievable within the course when the team prioritizes the service center, attendance confirmation, contextual resources, and preparation tasks. "
            "User testing will be essential to demonstrate that multimedia leaders and volunteers can find current information faster than in the existing WhatsApp-based process."
        ),
        "reflection_en": (
            "The rubric showed me that a good idea is not enough; every claim must be supported by objectives, a realistic method, and evidence. "
            "My multimedia experience gives me concrete cases to validate, especially late presentations, lyrics, sermon files, and last-minute changes. "
            "I need to turn those cases into clear user tasks instead of assuming that the interface will be intuitive. "
            "I will contribute by reviewing the multimedia workflow, preparing realistic demo data, checking responsive behavior, and recording usability findings in a consistent format."
        ),
    },
    {
        "first": "Felipe",
        "last": "Arce",
        "folder": "Arce_Felipe",
        "filename": "Arce_Felipe_1.3_APT122_AutoevaluacionFase1.docx",
        "focus_es": (
            "Mis intereses profesionales incluyen el desarrollo de soluciones digitales, el trabajo colaborativo y la mejora de experiencias mediante interfaces bien diseñadas. "
            "Mi servicio en multimedia aporta contacto directo con la preparación de contenidos y con la necesidad de coordinar información entre equipos. "
            "RUAHTONE me permite aplicar esos intereses en un proyecto con usuarios reales, reglas de negocio definidas y un resultado demostrable durante el periodo académico."
        ),
        "role_es": (
            "Aporto la segunda perspectiva de multimedia, apoyo en prototipado e implementación de componentes, consistencia visual y ejecución de pruebas funcionales y de usabilidad."
        ),
        "conclusion_en": (
            "The project definition establishes a practical path from a real coordination problem to a demonstrable software product. "
            "RUAHTONE is aligned with the graduate profile because the team must plan the project, model scalable information, build an integrated solution, and validate both the product and the process. "
            "The proposal remains feasible by limiting the first version to a local frontend with deterministic demo data. "
            "Success will depend on disciplined collaboration, traceable evidence, and early validation with people who serve in worship, sound, multimedia, and pastoral contexts."
        ),
        "reflection_en": (
            "This assessment made the relationship between our course competencies and the project more explicit. "
            "I understood that visual quality must support operational clarity and cannot replace functional evidence. "
            "Our main improvement area is to assign responsibilities and maintain a weekly record of decisions, tests, and unresolved risks. "
            "I will contribute by implementing reusable interface components, verifying complete user flows, supporting multimedia validation, and reporting defects with enough detail for the team to reproduce and correct them."
        ),
    },
]


OBJECTIVES = [
    (
        "Objetivo general",
        "Diseñar, construir y validar durante la asignatura un MVP frontend local de RUAHTONE que centralice la preparación de servicios de iglesias cristianas, permitiendo gestionar eventos, equipos, asignaciones, confirmaciones, repertorio, recursos y tareas desde una experiencia diferenciada por rol.",
    ),
    (
        "Objetivo específico 1",
        "Levantar y formalizar requisitos, reglas de negocio, roles, flujos y criterios de aceptación a partir de la experiencia del equipo y de validaciones con servidores de iglesias en Chile.",
    ),
    (
        "Objetivo específico 2",
        "Modelar una arquitectura y un conjunto de datos coherentes que separen personas, ministerios, servicios, asignaciones, canciones contextuales, recursos, checklists y actividad.",
    ),
    (
        "Objetivo específico 3",
        "Implementar un flujo vertical responsive y accesible para consultar el próximo servicio, confirmar asistencia, gestionar repertorio y recursos, completar tareas y observar la trazabilidad de los cambios.",
    ),
    (
        "Objetivo específico 4",
        "Evaluar el producto mediante pruebas automatizadas, recorridos funcionales y sesiones moderadas con usuarios representativos, utilizando los hallazgos para priorizar mejoras.",
    ),
]

PLAN = [
    ("1. Definición", "Semanas 1-3", "Problema, alcance, requisitos y rúbrica", "Equipo completo", "Experiencia en tres iglesias", "Alcance excesivo o supuestos sin validar"),
    ("2. Diseño", "Semanas 4-5", "Flujos, modelo, arquitectura y prototipo", "Equipo completo", "Documentación aprobada", "Diferencias entre procesos de cada iglesia"),
    ("3. Fundaciones", "Semanas 6-7", "Proyecto ejecutable, datos demo y navegación", "Desarrollo compartido", "Stack conocido y repositorio Git", "Curva de aprendizaje y coordinación técnica"),
    ("4. Flujo vertical", "Semanas 8-10", "Servicio, asignación, confirmación y actividad", "Desarrollo compartido", "Criterios de aceptación definidos", "Integración entre componentes y estado"),
    ("5. Preparación", "Semanas 11-12", "Repertorio, recursos y checklists", "Desarrollo compartido", "Modelo contextual de servicio", "Gestión de archivos solo simulada"),
    ("6. Calidad", "Semanas 13-14", "Pruebas, accesibilidad y correcciones", "Equipo completo", "Seed determinista y estrategia de QA", "Tiempo limitado para corregir hallazgos"),
    ("7. Validación", "Semanas 15-16", "Sesiones con usuarios, resultados y presentación", "Equipo completo", "Acceso a voluntarios reales", "Disponibilidad de participantes"),
]

EVIDENCE = [
    ("Especificación del producto", "Visión, requisitos, dominio, UX, arquitectura y decisiones", "Demuestra definición disciplinar y trazabilidad antes de construir."),
    ("Repositorio GitHub", "Commits, ramas, revisiones y documentación versionada", "Demuestra trabajo colaborativo, control de cambios y evolución del producto."),
    ("Prototipo/MVP ejecutable", "Flujo vertical persistente y responsive", "Evidencia la integración de componentes y el cumplimiento funcional."),
    ("Plan y tablero de trabajo", "Actividades, responsables, hitos, riesgos y estado", "Permite controlar el avance y justificar decisiones de alcance."),
    ("Pruebas y reportes", "Casos automatizados, checklist manual e incidencias", "Demuestra validación del producto y mejoras basadas en resultados."),
    ("Registro de validación", "Tareas, tiempos, éxito, dudas y comentarios de usuarios", "Permite contrastar la propuesta de valor con personas representativas."),
    ("Presentación final", "Problema, solución, demostración, resultados y próximos pasos", "Comunica de forma verificable el resultado y el aprendizaje del proyecto."),
]

COMPETENCIES = [
    ("Pruebas y certificación", "1.1 y 1.3", "Diseñar pruebas de validación y convertir sus resultados en mejoras priorizadas para el MVP."),
    ("Gestión de proyectos informáticos", "2.1 y 2.2", "Planificar alcance, actividades, recursos, responsables, riesgos e hitos; controlar el avance con evidencia versionada."),
    ("Modelos de datos escalables", "3.1 y 3.2", "Diseñar e implementar entidades y relaciones que soporten varias iglesias en el futuro sin mezclar la información de cada organización."),
    ("Desarrollo de soluciones", "4.1, 4.2 y 4.3", "Construir, integrar y desplegar una solución frontend coherente con técnicas sistemáticas de desarrollo y mantenimiento."),
]

SELF_ASSESSMENT = [
    (1, "Descripción y relevancia laboral", "Completamente logrado", "Se explica el problema, la solución, el mercado inicial y su relación con el desarrollo de software."),
    (2, "Competencias del perfil de egreso", "Completamente logrado", "Se relacionan pruebas, gestión, datos y desarrollo con actividades concretas del proyecto."),
    (3, "Intereses profesionales", "Completamente logrado", "Se explicitan intereses personales, experiencia e influencia en el aporte individual."),
    (4, "Factibilidad", "Completamente logrado", "Se delimitan tiempo, materiales, factores externos, riesgos y medidas de mitigación."),
    (5, "Objetivos", "Completamente logrado", "Se formula un objetivo general y cuatro objetivos específicos claros y verificables."),
    (6, "Metodología", "Completamente logrado", "Se propone trabajo incremental, control de versiones, pruebas y validación con usuarios."),
    (7, "Plan de trabajo", "Completamente logrado", "El plan contiene actividades, duración, responsables, recursos, facilitadores y obstaculizadores."),
    (8, "Evidencias", "Completamente logrado", "Cada evidencia se vincula con el logro de actividades y resultados del proyecto."),
    (9, "Redacción, citas y referencias", "Completamente logrado", "El informe fue revisado y utiliza referencias internas en formato APA 7 adaptado a documentos del proyecto."),
    (10, "Formato disciplinar", "Completamente logrado", "Incluye identificación, abstracts, desarrollo, conclusiones, reflexión, matriz y referencias."),
    (11, "Indicadores de calidad", "Completamente logrado", "La propuesta cubre los indicadores seleccionados del perfil de egreso con evidencia planificada."),
    (12, "Inglés intermedio alto", "Completamente logrado", "Abstract, conclusion and reflection use connected ideas and topic-specific vocabulary."),
]


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=11, color=INK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def configure_document(student):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = rgb(BLUE)
    title.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name, left, first in (
        ("List Bullet", 0.375, -0.194),
        ("List Number", 0.375, -0.194),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(left)
        style.paragraph_format.first_line_indent = Inches(first)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Table Citation" not in [s.name for s in styles]:
        citation = styles.add_style("Table Citation", WD_STYLE_TYPE.PARAGRAPH)
        citation.font.name = "Calibri"
        citation.font.size = Pt(9)
        citation.font.italic = True
        citation.font.color.rgb = rgb(MUTED)
        citation.paragraph_format.space_before = Pt(4)
        citation.paragraph_format.space_after = Pt(4)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("RUAHTONE  /  AUTOEVALUACIÓN DEFINICIÓN PROYECTO APT")
    set_font(hr, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run(f"{student['first']} {student['last']}  •  {COURSE_CODE}  •  Sección {SECTION}")
    set_font(fr, size=8.5, color=MUTED)
    return doc


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("  •  Página ")
    set_font(run, size=8.5, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, txt, fld_end])


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, size=9, color=WHITE, bold=True)
    set_repeat_header(table.rows[0])
    for row_number, row_data in enumerate(rows):
        cells = table.add_row().cells
        set_cant_split(table.rows[-1])
        for index, value in enumerate(row_data):
            if row_number % 2:
                set_cell_shading(cells[index], "FAFBFD")
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_font(r, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph(style="Table Citation")
    return table


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360])
    # Layout callouts have a single identifying row. Mark it explicitly so
    # assistive tooling does not report the table as structurally headerless.
    set_repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label.upper())
    set_font(r, size=9, color=accent, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=10.5)
    doc.add_paragraph(style="Table Citation")


def add_cover(doc, student):
    for _ in range(2):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    kr = kicker.add_run("EVALUACIÓN FASE 1  ·  AUTOEVALUACIÓN")
    set_font(kr, size=10, color=DARK_BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    tr = title.add_run("Definición Proyecto APT")
    set_font(tr, size=28, color=BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    sr = subtitle.add_run("RUAHTONE")
    set_font(sr, size=18, color=INK, bold=True)

    add_callout(
        doc,
        "Propósito del informe",
        "Presentar y autoevaluar la definición del Proyecto APT mediante los doce indicadores de la pauta, diferenciando el aporte individual del trabajo grupal.",
    )

    metadata = [
        ("Estudiante", f"{student['first']} {student['last']}"),
        ("Equipo", "Benjamín Olmedo · Daniel Baeza · Felipe Arce"),
        ("Asignatura", f"{COURSE_CODE} · {COURSE}"),
        ("Sección", SECTION),
        ("Proyecto", PROJECT),
        ("Fecha", DATE_TEXT),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in metadata:
        cells = table.add_row().cells
        set_cell_shading(cells[0], LIGHT_GRAY)
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_font(r, size=10, color=DARK_BLUE, bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=10)
    # The metadata grid has no separate column-heading row; its first row is
    # the identifying row for accessibility and document-audit purposes.
    set_repeat_header(table.rows[0])
    set_table_geometry(table, [2200, 7160])
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_paragraph(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)


def build_document(student):
    doc = configure_document(student)
    add_page_number(doc.sections[0].footer.paragraphs[0])
    add_cover(doc, student)

    doc.add_heading("Resumen / Abstract", level=1)
    doc.add_heading("Resumen en español", level=2)
    add_paragraph(
        doc,
        "RUAHTONE es una propuesta de aplicación web y, en su proyección, móvil para organizar servicios y eventos de iglesias cristianas. Surge ante la fragmentación de horarios, confirmaciones, repertorios, tonalidades, letras, prédicas y recursos multimedia entre WhatsApp, llamadas y archivos dispersos. La solución centraliza cada servicio y relaciona personas, ministerios, funciones, canciones, recursos, tareas y actividad, mediante permisos escalables para servidores, líderes y administradores. Durante la asignatura se construirá un MVP frontend local, responsive y persistente en el navegador, sin backend, pagos, autenticación real ni integración con WhatsApp. La metodología será incremental, con requisitos versionados, arquitectura por capas, control de versiones, pruebas y validación con usuarios de iglesias en Chile. La evidencia permitirá evaluar si las personas encuentran y actualizan la información crítica con mayor claridad que en el proceso actual."
    )
    p = doc.add_paragraph()
    r = p.add_run("Palabras clave: ")
    set_font(r, bold=True)
    r = p.add_run("gestión de eventos, iglesias, coordinación de equipos, software, experiencia de usuario.")
    set_font(r, italic=True)

    doc.add_heading("Abstract in English", level=2)
    add_paragraph(
        doc,
        "RUAHTONE is a web-based and, in its long-term vision, mobile application for organizing services and events in Christian churches. The project responds to fragmented schedules, attendance confirmations, song sets, key changes, lyrics, sermon files, and multimedia resources distributed across WhatsApp, phone calls, and isolated documents. The solution places each service at the center and connects people, ministries, assignments, songs, resources, checklists, and structured activity through scalable permissions for volunteers, leaders, and administrators. During the course, the team will build a responsive local frontend MVP with browser persistence, excluding backend services, real authentication, payments, and WhatsApp integration. The project will follow an incremental method supported by versioned requirements, layered architecture, source control, testing, and validation with church users in Chile. The collected evidence will determine whether users can find and update critical information more clearly than in the current process."
    )
    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    set_font(r, bold=True)
    r = p.add_run("event management, churches, team coordination, software, user experience.")
    set_font(r, italic=True)

    doc.add_heading("1. Descripción y relevancia del Proyecto APT", level=1)
    add_paragraph(
        doc,
        "RUAHTONE propone centralizar la planificación de cultos dominicales, reuniones de jóvenes, ensayos, reuniones de planificación y otros eventos. Actualmente, muchas iglesias coordinan por WhatsApp y llamadas; esto provoca información cruzada, contenidos que llegan tarde, cambios de tonalidad difíciles de seguir y poca claridad sobre horarios, lugares, funciones y confirmaciones. El producto organiza cada evento como un centro operacional con equipos de sonido, multimedia, adoración y pastorado, además de funciones específicas por área."
    )
    add_paragraph(
        doc,
        "El sistema considera tres niveles organizacionales acumulativos: el servidor opera únicamente sus responsabilidades; el líder conserva las capacidades de servidor y coordina su área; y el administrador conserva capacidades de liderazgo y gobierna la organización. La visión comercial futura contempla suscripciones escalables y soporte para múltiples sedes en el plan superior, pero estos elementos no forman parte del MVP académico."
    )
    add_paragraph(
        doc,
        "La propuesta es relevante para el campo laboral de la informática porque integra descubrimiento de producto, análisis de requisitos, arquitectura, modelado de datos, interfaces responsive, accesibilidad, persistencia, control de versiones y pruebas. También aborda una oportunidad concreta: desarrollar software vertical para organizaciones que todavía dependen de procesos manuales y canales no diseñados para gestión operacional (RUAHTONE, 2026a, 2026b)."
    )
    add_callout(
        doc,
        "Alcance evaluable",
        "El proyecto de la asignatura es un MVP frontend local que valida el núcleo de organización. Las aplicaciones nativas, el SaaS multiiglesia, las suscripciones, las notificaciones reales y WhatsApp quedan como evolución posterior.",
        fill=LIGHT_GRAY,
        accent=DARK_BLUE,
    )

    doc.add_heading("2. Relación con el perfil de egreso", level=1)
    add_paragraph(
        doc,
        "Las competencias seleccionadas se aplican de forma integrada. La tabla identifica los indicadores de calidad disciplinarios pertinentes y la manera en que el proyecto generará evidencia de su logro."
    )
    add_table(doc, ["Competencia", "Indicadores", "Aplicación en RUAHTONE"], COMPETENCIES, [2550, 1350, 5460], font_size=9.1)

    doc.add_heading("3. Relación con mis intereses profesionales", level=1)
    add_paragraph(doc, student["focus_es"])
    add_paragraph(doc, "Mi aporte individual: " + student["role_es"])

    doc.add_heading("4. Factibilidad", level=1)
    add_paragraph(
        doc,
        "El proyecto es factible si se mantiene el alcance acordado. El equipo dispone de tres estudiantes con experiencia en iglesias distintas, acceso a usuarios representativos y un repositorio GitHub para colaborar. El MVP utiliza tecnologías web gratuitas y datos ficticios, por lo que no requiere infraestructura pagada ni tratamiento de información sensible. El trabajo se divide en incrementos verticales que producen resultados demostrables durante el semestre (RUAHTONE, 2026e)."
    )
    add_table(
        doc,
        ["Dimensión", "Condición y respuesta"],
        [
            ("Tiempo", "Plan semestral por fases; priorización del flujo central antes de módulos secundarios."),
            ("Materiales", "Tres computadores, conexión a internet para GitHub, navegadores y herramientas de desarrollo gratuitas."),
            ("Factores externos", "Disponibilidad de voluntarios para validación y retroalimentación del docente."),
            ("Riesgo de alcance", "Congelar backend, pagos, aplicaciones nativas, WhatsApp y multi-sede durante el MVP."),
            ("Riesgo técnico", "Arquitectura por capas, datos demo versionados, revisión por pares y pruebas tempranas."),
            ("Riesgo de coordinación", "Ramas, pull requests pequeños, responsables visibles y reuniones breves de seguimiento."),
        ],
        [2000, 7360],
        font_size=9.4,
    )

    doc.add_heading("5. Objetivos", level=1)
    for label, text in OBJECTIVES:
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = True
        r = p.add_run(label + ". ")
        set_font(r, color=DARK_BLUE, bold=True)
        r = p.add_run(text)
        set_font(r)

    doc.add_heading("6. Metodología de trabajo", level=1)
    add_paragraph(
        doc,
        "Se utilizará una metodología incremental y colaborativa, adecuada a un producto digital en fase de descubrimiento. Cada ciclo comenzará con un problema y criterios de aceptación, continuará con diseño e implementación, y terminará con verificación, evidencia y ajuste."
    )
    add_numbered(
        doc,
        [
            "Descubrir y validar necesidades mediante entrevistas breves y observación de procesos en las iglesias del equipo.",
            "Especificar historias, requisitos, reglas de negocio, flujos, riesgos y criterios de aceptación en documentos versionados.",
            "Diseñar el modelo de datos, la arquitectura, la navegación y prototipos antes de implementar funcionalidades de alto riesgo.",
            "Construir incrementos verticales pequeños en ramas separadas, revisados mediante pull request antes de integrarlos a main.",
            "Verificar dominio, componentes y flujos críticos; revisar accesibilidad, responsive, persistencia y consistencia de datos.",
            "Validar con usuarios representativos, registrar tiempos, éxito, dudas y errores, y priorizar mejoras por impacto y frecuencia.",
        ],
    )

    doc.add_heading("7. Plan de trabajo", level=1)
    add_table(
        doc,
        ["Actividad", "Duración", "Producto", "Responsable", "Facilitador", "Obstáculo"],
        PLAN,
        [1250, 1150, 1900, 1400, 1750, 1910],
        font_size=8.1,
    )
    add_paragraph(
        doc,
        "Recursos transversales: computadores personales, GitHub, entorno de desarrollo web, documentación del proyecto, datos demo, instrumentos de prueba y acceso programado a participantes. El calendario exacto se ajustará al cronograma oficial de la asignatura sin cambiar el orden de dependencias."
    )

    doc.add_heading("8. Evidencias de logro", level=1)
    add_table(doc, ["Evidencia", "Contenido", "Justificación"], EVIDENCE, [2200, 3100, 4060], font_size=8.8)

    doc.add_heading("9. Indicadores de calidad disciplinarios", level=1)
    add_paragraph(
        doc,
        "La definición incorpora los indicadores 1.1, 1.3, 2.1, 2.2, 3.1, 3.2, 4.1, 4.2 y 4.3. En esta fase se evidencia principalmente el diseño de pruebas, la planificación, el modelo escalable y la arquitectura de integración. La aplicación, el control y la implantación se demostrarán progresivamente mediante el repositorio, el MVP y los reportes de validación."
    )
    add_bullets(
        doc,
        [
            "Pruebas: estrategia de dominio, componentes, recorridos end-to-end y verificación manual (RUAHTONE, 2026f).",
            "Gestión: roadmap por incrementos, definición de terminado, GitHub y seguimiento de riesgos (RUAHTONE, 2026e).",
            "Datos: entidades normalizadas, relaciones contextuales, invariantes y separación futura por iglesia (RUAHTONE, 2026c).",
            "Desarrollo: capas UI, aplicación, dominio e infraestructura, integradas mediante acciones y repositorios (RUAHTONE, 2026d).",
        ],
    )

    doc.add_heading("10. Autoevaluación según la rúbrica", level=1)
    add_callout(
        doc,
        "Resultado declarado",
        "Nivel objetivo: Completamente logrado. Esta declaración se fundamenta en la evidencia indicada y debe ser contrastada con la retroalimentación del docente.",
        fill="EFF8F3",
        accent=GREEN,
    )
    add_table(doc, ["N.º", "Indicador", "Nivel", "Evidencia en el informe"], SELF_ASSESSMENT, [520, 2500, 1750, 4590], font_size=8.5)

    doc.add_heading("11. Conclusion", level=1)
    add_paragraph(doc, student["conclusion_en"])

    doc.add_heading("12. Personal reflection", level=1)
    add_paragraph(doc, student["reflection_en"])

    doc.add_page_break()
    references_heading = doc.add_heading("Referencias", level=1)
    references_heading.paragraph_format.space_before = Pt(6)
    references_heading.paragraph_format.space_after = Pt(3)
    references = [
        "RUAHTONE. (2026a). Visión y alcance [Documento interno]. Repositorio RUAHTONE.",
        "RUAHTONE. (2026b). Requisitos del producto [Documento interno]. Repositorio RUAHTONE.",
        "RUAHTONE. (2026c). Modelo de dominio [Documento interno]. Repositorio RUAHTONE.",
        "RUAHTONE. (2026d). Arquitectura técnica [Documento interno]. Repositorio RUAHTONE.",
        "RUAHTONE. (2026e). Roadmap del MVP [Documento interno]. Repositorio RUAHTONE.",
        "RUAHTONE. (2026f). Estrategia de calidad [Documento interno]. Repositorio RUAHTONE.",
    ]
    for reference in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(reference)
        set_font(r, size=9)

    props = doc.core_properties
    props.title = f"Autoevaluación Definición Proyecto APT - {student['first']} {student['last']}"
    props.subject = "Evaluación Fase 1 - RUAHTONE"
    props.author = f"{student['first']} {student['last']}"
    props.keywords = "Capstone, Proyecto APT, RUAHTONE, autoevaluación, fase 1"
    props.comments = "Documento individual construido desde la definición oficial del proyecto RUAHTONE."

    target_dir = OUTPUT_ROOT / student["folder"]
    target_dir.mkdir(parents=True, exist_ok=True)
    output_path = target_dir / student["filename"]
    doc.save(output_path)
    return output_path


def main():
    outputs = [build_document(student) for student in STUDENTS]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
